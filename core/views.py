from django.shortcuts import render, redirect, get_object_or_404
from django.http import JsonResponse
from .models import Classe, Eleve, Periode, Matiere,Notation
from .forms import NotationForm,RegisterForm, UpdateProfileForm
from django.contrib import messages
from django.contrib.auth import authenticate, login,logout
from django.contrib.auth.decorators import login_required
from django.db.models import Avg
from django.db.models import Sum


def signup(request):
    form = RegisterForm()
    if request.method == 'POST':
        form = RegisterForm(request.POST)
        if form.is_valid():
            form.save()
            messages.success(request, "Compte créé avec succès.")
            return redirect("signin")  # Redirect to signin page after successful signup
        
    context = {"form": form}
    return render(request, "core/signup.html", context)

def signin(request):
    if request.method == 'POST':
        email = request.POST["email"]
        password = request.POST["password"]

        user = authenticate(request, email=email, password=password)  # Utilisation de 'username' avec l'email
        if user is not None:
            login(request, user)
            return redirect('index')
        else:
            messages.error(request, "Identifiants invalides. Veuillez réessayer.")

    return render(request, "core/login.html")

def signout(request):
    logout(request)
    return redirect("index")

@login_required(login_url="signin")
def profile(request):
    user = request.user
    
    context={"user": user, }
    return render(request, "core/profile.html", context)

@login_required(login_url="signin")
def update_profile(request):
    if request.user.is_authenticated:
        user = request.user
        form = UpdateProfileForm(instance=user)
        if request.method == 'POST':
            form = UpdateProfileForm(request.POST, request.FILES, instance=user)
            if form.is_valid():
                form.save()
                messages.success(request, "Profile updated successfully")
                return redirect("profile")
                
    context = {"form": form}
    return render(request, "core/update_profile.html", context)

@login_required(login_url="signin")
def update_profile(request):
    if request.user.is_authenticated:
        user = request.user
        form = UpdateProfileForm(instance=user)
        if request.method == 'POST':
            form = UpdateProfileForm(request.POST, request.FILES, instance=user)
            if form.is_valid():
                form.save()
                messages.success(request, "Profile updated successfully")
                return redirect("profile")
                
    context = {"form": form}
    return render(request, "core/update_profile.html", context)


def ajouter_notation(request):
    if request.method == 'POST':
        eleve_id = request.POST.get('eleve')
        matiere_id = request.POST.get('matiere')
        periode_id = request.POST.get('periode')

        # Récupérer ou créer la notation
        notation, created = Notation.objects.get_or_create(
            eleve_id=eleve_id,
            matiere_id=matiere_id,
            periode_id=periode_id,
            defaults={'note_attendue': 0, 'note_obtenue': 0}
        )

        # Remplir le formulaire avec les valeurs existantes
        form = NotationForm(request.POST, instance=notation)
        if form.is_valid():
            form.save()
            messages.success(request, "Notation ajoutée avec succès !")  # Affichage d'un message de succès
            return redirect('ajouter_notation')  # Remplace par ton URL de succès
    else:
        classe_id = request.GET.get('classe_id')
        form = NotationForm()  # Initialisez sans instance

    classes = Classe.objects.all()
    periodes = Periode.objects.all()

    return render(request, 'core/ajouter_notation.html', {'form': form, 'classes': classes, 'periodes': periodes})



def load_eleves(request):
    classe_id = request.GET.get('classe_id')
    eleves = Eleve.objects.filter(classe_id=classe_id).values('id', 'nom')
    return render(request, 'core/eleve_dropdown_list_options.html', {'eleves': eleves})

def load_matieres(request):
    classe_id = request.GET.get('classe_id')
    matieres = Matiere.objects.filter(classe_id=classe_id).values('id', 'nom')
    return render(request, 'core/matiere_dropdown_list_options.html', {'matieres': matieres})



def success_view(request):
    return render(request, 'core/success.html')  # Assurez-vous que ce template existe

# Liste des classes

def liste_classes(request):
    search_query = request.GET.get('search', '')  # Récupère la requête de recherche depuis l'URL
    
    # Filtrer les classes par nom si la requête de recherche existe
    if search_query:
        classes = Classe.objects.filter(nom__icontains=search_query)  # Filtrage insensible à la casse
    else:
        classes = Classe.objects.all()
    
    return render(request, 'core/liste_classes.html', {'classes': classes, 'search_query': search_query})
# Détails d'une classe
def details_classe(request, classe_id):
    classe = get_object_or_404(Classe, id=classe_id)
    eleves = Eleve.objects.filter(classe=classe)
    matieres = Matiere.objects.filter(classe=classe)
    return render(request, 'core/details_classe.html', {'classe': classe, 'eleves': eleves, 'matieres': matieres})

# Détails d'un élève
# def details_eleve(request, eleve_id):
#     eleve = get_object_or_404(Eleve, id=eleve_id)
#     notations = Notation.objects.filter(eleve=eleve)
#     return render(request, 'core/details_eleve.html', {'eleve': eleve, 'notations': notations})

# from django.shortcuts import render, get_object_or_404
# from .models import Eleve, Notation, Periode

# def details_eleve(request, eleve_id):
#     eleve = get_object_or_404(Eleve, id=eleve_id)
#     periodes = Periode.objects.all()
#     matieres = Matiere.objects.filter(classe=eleve.classe)

#     notations_par_matiere = {matiere: {} for matiere in matieres}

#     for periode in periodes:
#         notations = Notation.objects.filter(eleve=eleve, periode=periode)
#         for notation in notations:
#             notations_par_matiere[notation.matiere][periode] = notation

#     return render(request, 'core/details_eleve.html', {
#         'eleve': eleve,
#         'periodes': periodes,
#         'notations_par_matiere': notations_par_matiere,
#     })

def details_eleve(request, eleve_id):
    # Récupérer l'élève à partir de l'ID
    eleve = get_object_or_404(Eleve, id=eleve_id)
    
    # Récupérer toutes les périodes disponibles
    periodes = Periode.objects.all()
    
    notations_par_periode = {}
    
    for periode in periodes:
        # Filtrer les notations de l'élève pour chaque période
        notations = Notation.objects.filter(eleve=eleve, periode=periode)
        
        if notations.exists():
            # Calcul du total des notes attendues et obtenues pour chaque période
            total_attendu = notations.aggregate(Sum('note_attendue'))['note_attendue__sum'] or 0
            total_obtenu = notations.aggregate(Sum('note_obtenue'))['note_obtenue__sum'] or 0
            
            # Calcul du pourcentage
            pourcentage = (total_obtenu / total_attendu * 100) if total_attendu > 0 else 0
            
            notations_par_periode[periode] = {
                'notations': notations,
                'pourcentage': pourcentage,
            }

    # Renvoyer le rendu avec les données nécessaires
    return render(request, 'core/details_eleve.html', {
        'eleve': eleve,
        'notations_par_periode': notations_par_periode,
        'periodes': periodes,
    })

def details_matiere(request, matiere_id):
    matiere = get_object_or_404(Matiere, id=matiere_id)
    notations = Notation.objects.filter(matiere=matiere).select_related('eleve', 'periode')

    # Regrouper les notations par période
    periodes_notes = {}
    for notation in notations:
        periode = notation.periode
        if periode not in periodes_notes:
            periodes_notes[periode] = []
        eleve = notation.eleve
        periodes_notes[periode].append({
            'eleve': eleve,
            'note_attendue': notation.note_attendue,
            'note_obtenue': notation.note_obtenue,
        })

    return render(request, 'core/details_matiere.html', {
        'matiere': matiere,
        'prof': matiere.prof,  # Ajouter l'objet prof pour l'afficher dans le template
        'periodes_notes': periodes_notes,
    })


def details_periode_matiere(request, matiere_id, periode_id):
    matiere = get_object_or_404(Matiere, id=matiere_id)
    periode = get_object_or_404(Periode, id=periode_id)
    
    notations = Notation.objects.filter(matiere=matiere, periode=periode).select_related('eleve')
    
    return render(request, 'core/details_periode_matiere.html', {
        'matiere': matiere,
        'periode': periode,
        'prof': matiere.prof,
        'notations': notations,
    })


def details_periode(request, eleve_id, periode_id):
    eleve = get_object_or_404(Eleve, id=eleve_id)
    periode = get_object_or_404(Periode, id=periode_id)
    notations = Notation.objects.filter(eleve=eleve, periode=periode)

    total_attendu = notations.aggregate(Sum('note_attendue'))['note_attendue__sum'] or 0
    total_obtenu = notations.aggregate(Sum('note_obtenue'))['note_obtenue__sum'] or 0

    # Calcul du pourcentage
    pourcentage = (total_obtenu / total_attendu * 100) if total_attendu > 0 else 0

    return render(request, 'core/details_periode.html', {
        'eleve': eleve,
        'periode': periode,
        'notations': notations,
        'total_attendu': total_attendu,
        'total_obtenu': total_obtenu,
        'pourcentage': pourcentage,
    })



import openpyxl
from django.http import HttpResponse
from io import BytesIO
from django.shortcuts import get_object_or_404

def generer_excel(request, periode_id):
    periode = get_object_or_404(Periode, id=periode_id)
    
    # Récupération des notations avec l'élève et la matière
    notations = Notation.objects.select_related('eleve', 'matiere').filter(periode=periode)
    
    if notations.exists():
        classe = notations.first().eleve.classe.nom
    else:
        return HttpResponse("Aucune notation disponible pour cette période.", status=404)

    # Création du fichier Excel
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = f"Rapport {periode.nom}"

    # Ajout du nom de la classe
    ws.append([f"Classe: {classe}"])
    ws.append([])  # Ligne vide

    # Dictionnaire pour stocker les totaux par élève et matière
    totals = {}

    for notation in notations:
        eleve_nom = notation.eleve.nom
        matiere_nom = notation.matiere.nom
        note_attendue = notation.note_attendue
        note_obtenue = notation.note_obtenue

        # Structure des totaux
        if eleve_nom not in totals:
            totals[eleve_nom] = {}
        
        if matiere_nom not in totals[eleve_nom]:
            totals[eleve_nom][matiere_nom] = {'attendue': 0, 'obtenue': 0}
        
        totals[eleve_nom][matiere_nom]['attendue'] += note_attendue
        totals[eleve_nom][matiere_nom]['obtenue'] += note_obtenue

    # Ajout du nom de chaque élève dans le fichier
    for eleve in totals.keys():
        ws.append([f"Élève: {eleve}"])
    
    # Ajout d'une ligne vide pour séparer les élèves des en-têtes
    ws.append([])

    # Ajout des en-têtes
    ws.append(["Élève", "Matière", "Total Note Attendue", "Total Note Obtenue"])

    # Remplissage du fichier avec les données
    for eleve, matieres in totals.items():
        for matiere, notes in matieres.items():
            total_attendue = notes['attendue']
            total_obtenue = notes['obtenue']
            ws.append([eleve, matiere, total_attendue, total_obtenue])

    # Enregistrement du fichier en mémoire
    output = BytesIO()
    wb.save(output)
    output.seek(0)

    # Création de la réponse HTTP pour le fichier
    response = HttpResponse(output, content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')
    response['Content-Disposition'] = f'attachment; filename="rapport_{periode.nom}.xlsx"'
    
    return response


from django.http import HttpResponse 
from openpyxl import Workbook
from django.shortcuts import get_object_or_404
from .models import Eleve, Notation, Periode, Matiere

def generer_excel2(request, eleve_id):
    eleve = get_object_or_404(Eleve, id=eleve_id)
    periodes = Periode.objects.all()  # Récupérer toutes les périodes
    matieres = Matiere.objects.filter(classe=eleve.classe)  # Récupérer les matières de la classe de l'élève

    # Créer un classeur et une feuille
    wb = Workbook()
    ws = wb.active
    ws.title = f"Situation de {eleve.nom}"

    # Ajouter les en-têtes
    headers = ['Matières']
    for periode in periodes:
        headers.append(f"{periode.nom} - Note Obtenue")
        headers.append(f"{periode.nom} - Note Attendue")
    headers.append('Total Obtenu')
    headers.append('Total Attendu')
    for periode in periodes:
        headers.append(f"{periode.nom} - Pourcentage")
    ws.append(headers)

    # Variables pour accumuler les totaux
    total_obtenu_periode = [0] * len(periodes)
    total_attendu_periode = [0] * len(periodes)

    # Ajouter les données des notations
    for matiere in matieres:
        row = [matiere.nom]  # Commencer la ligne avec le nom de la matière
        for index, periode in enumerate(periodes):
            notation = Notation.objects.filter(eleve=eleve, matiere=matiere, periode=periode).first()
            if notation:
                row.append(notation.note_obtenue)
                row.append(notation.note_attendue)
                total_obtenu_periode[index] += notation.note_obtenue
                total_attendu_periode[index] += notation.note_attendue
            else:
                row.append('N/A')
                row.append('N/A')
        ws.append(row)

    # Ajouter la ligne de totaux
    total_row = ['Total']
    for total_obtenu, total_attendu in zip(total_obtenu_periode, total_attendu_periode):
        total_row.append(total_obtenu)
        total_row.append(total_attendu)
    ws.append(total_row)

    # Calculer et ajouter les pourcentages
    pourcentage_row = ['Pourcentage']
    for total_obtenu, total_attendu in zip(total_obtenu_periode, total_attendu_periode):
        if total_attendu > 0:
            pourcentage = (total_obtenu / total_attendu) * 100
        else:
            pourcentage = 0
        pourcentage_row.append(f"{pourcentage:.2f}%")
    ws.append(pourcentage_row)

    # Créer la réponse HTTP
    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')
    response['Content-Disposition'] = f'attachment; filename="{eleve.nom}_situation.xlsx"'

    # Sauvegarder le classeur dans la réponse
    wb.save(response)
    
    return response

def situations_eleve(request, eleve_id):
    eleve = get_object_or_404(Eleve, id=eleve_id)
    periodes = Periode.objects.all()
    matieres = Matiere.objects.filter(classe=eleve.classe)

    notations_par_matiere = {
        matiere: {periode: Notation.objects.filter(eleve=eleve, matiere=matiere, periode=periode).first() for periode in periodes}
        for matiere in matieres
    }

    total_obtenu = 0
    total_attendu = 0

    total_periode_obtenu = {periode: 0 for periode in periodes}
    total_periode_attendu = {periode: 0 for periode in periodes}

    for matiere, notations in notations_par_matiere.items():
        for periode, notation in notations.items():
            if notation:
                total_obtenu += notation.note_obtenue
                total_attendu += notation.note_attendue
                total_periode_obtenu[periode] += notation.note_obtenue
                total_periode_attendu[periode] += notation.note_attendue

    # Calculer le pourcentage par période
    pourcentage_periode = {}
    for periode in periodes:
        if total_periode_attendu[periode] > 0:
            pourcentage_periode[periode] = (total_periode_obtenu[periode] / total_periode_attendu[periode]) * 100
        else:
            pourcentage_periode[periode] = 0

    # Calculer le pourcentage total
    pourcentage_total = (total_obtenu / total_attendu * 100) if total_attendu > 0 else 0

    return render(request, 'core/situations_eleve.html', {
        'eleve': eleve,
        'periodes': periodes,
        'matieres': matieres,
        'notations_par_matiere': notations_par_matiere,
        'total_obtenu': total_obtenu,
        'total_attendu': total_attendu,
        'total_periode_obtenu': total_periode_obtenu,
        'total_periode_attendu': total_periode_attendu,
        'pourcentage_periode': pourcentage_periode,
        'pourcentage_total': pourcentage_total,
    })

from django.http import HttpResponse
from django.shortcuts import get_object_or_404
from django.views import View
from .models import Eleve, Periode, Matiere, Notation
from .utils import render_to_pdf

class GeneratePdf(View):
    def get(self, request, *args, **kwargs):
        eleve_id = kwargs.get('eleve_id')
        eleve = get_object_or_404(Eleve, id=eleve_id)
        
        periodes = Periode.objects.all()
        matieres = Matiere.objects.filter(classe=eleve.classe)

        # Récupération des notations
        notations_par_matiere = {
            matiere: {periode: Notation.objects.filter(eleve=eleve, matiere=matiere, periode=periode).first() 
                      for periode in periodes}
            for matiere in matieres
        }

        # Calcul des totaux
        total_obtenu = 0
        total_attendu = 0
        total_periode_obtenu = {periode: 0 for periode in periodes}
        total_periode_attendu = {periode: 0 for periode in periodes}

        for matiere, notations in notations_par_matiere.items():
            for periode, notation in notations.items():
                if notation:
                    total_obtenu += notation.note_obtenue
                    total_attendu += notation.note_attendue
                    total_periode_obtenu[periode] += notation.note_obtenue
                    total_periode_attendu[periode] += notation.note_attendue

        pourcentage_periode = {}
        for periode in periodes:
            if total_periode_attendu[periode] > 0:
                pourcentage_periode[periode] = (total_periode_obtenu[periode] / total_periode_attendu[periode]) * 100
            else:
                pourcentage_periode[periode] = 0

        pourcentage_total = (total_obtenu / total_attendu * 100) if total_attendu > 0 else 0

        context = {
            'eleve': eleve,
            'periodes': periodes,
            'matieres': matieres,
            'notations_par_matiere': notations_par_matiere,
            'total_obtenu': total_obtenu,
            'total_attendu': total_attendu,
            'total_periode_obtenu': total_periode_obtenu,
            'total_periode_attendu': total_periode_attendu,
            'pourcentage_periode': pourcentage_periode,
            'pourcentage_total': pourcentage_total,
        }

        # Génération du PDF
        pdf = render_to_pdf('core/situations_eleve.html', context)

        if pdf:
            # Renvoie le PDF avec un en-tête Content-Disposition pour le téléchargement
            response = HttpResponse(pdf, content_type='application/pdf')
            response['Content-Disposition'] = f'attachment; filename="rapport_{eleve.nom}_{eleve.prenom}.pdf"'
            return response
        else:
            return HttpResponse("Erreur lors de la génération du PDF", status=500)

# pdf situation eleve
# from weasyprint import HTML
# from django.http import HttpResponse
# from django.template.loader import render_to_string

# def generate_pdf(request, eleve_id):
#     # Récupération des données de l'élève
#     eleve = get_object_or_404(Eleve, id=eleve_id)
#     periodes = Periode.objects.all()
#     matieres = Matiere.objects.filter(classe=eleve.classe)

#     notations_par_matiere = {
#         matiere: {periode: Notation.objects.filter(eleve=eleve, matiere=matiere, periode=periode).first() for periode in periodes}
#         for matiere in matieres
#     }

#     total_obtenu = 0
#     total_attendu = 0

#     total_periode_obtenu = {periode: 0 for periode in periodes}
#     total_periode_attendu = {periode: 0 for periode in periodes}

#     for matiere, notations in notations_par_matiere.items():
#         for periode, notation in notations.items():
#             if notation:
#                 total_obtenu += notation.note_obtenue
#                 total_attendu += notation.note_attendue
#                 total_periode_obtenu[periode] += notation.note_obtenue
#                 total_periode_attendu[periode] += notation.note_attendue

#     # Calculer le pourcentage par période
#     pourcentage_periode = {}
#     for periode in periodes:
#         if total_periode_attendu[periode] > 0:
#             pourcentage_periode[periode] = (total_periode_obtenu[periode] / total_periode_attendu[periode]) * 100
#         else:
#             pourcentage_periode[periode] = 0

#     # Calculer le pourcentage total
#     pourcentage_total = (total_obtenu / total_attendu * 100) if total_attendu > 0 else 0

#     # Rendre le HTML à partir du template
#     html_string = render_to_string('core/situations_eleve.html', {
#         'eleve': eleve,
#         'periodes': periodes,
#         'matieres': matieres,
#         'notations_par_matiere': notations_par_matiere,
#         'total_obtenu': total_obtenu,
#         'total_attendu': total_attendu,
#         'total_periode_obtenu': total_periode_obtenu,
#         'total_periode_attendu': total_periode_attendu,
#         'pourcentage_periode': pourcentage_periode,
#         'pourcentage_total': pourcentage_total,
#     })

#     # Créer le PDF à partir du HTML
#     html = HTML(string=html_string)
#     pdf_file = html.write_pdf()

#     # Retourner le PDF comme une réponse HTTP
#     response = HttpResponse(pdf_file, content_type='application/pdf')
#     response['Content-Disposition'] = f'attachment; filename="situation_eleve_{eleve.nom}_{eleve.prenom}.pdf"'
#     return response
# views.py
from django.http import HttpResponse
from django.shortcuts import get_object_or_404
from .models import Eleve, Notation, Periode, Matiere
from docx import Document

def generer_word(request, eleve_id):
    eleve = get_object_or_404(Eleve, id=eleve_id)
    periodes = Periode.objects.all()  # Récupérer toutes les périodes
    matieres = Matiere.objects.filter(classe=eleve.classe)  # Récupérer les matières de la classe de l'élève

    # Vérification que les périodes et matières existent
    if not periodes:
        return HttpResponse("Aucune période trouvée pour cet élève.", status=404)
    if not matieres:
        return HttpResponse("Aucune matière trouvée pour cet élève.", status=404)

    # Créer un document Word
    doc = Document()

    # Ajouter le titre
    doc.add_heading(f'Situation de {eleve.nom}', 0)

    # Calculer le nombre de colonnes
    num_cols = 1 + len(periodes) * 2 + 2 + len(periodes)

    # Ajouter un tableau avec le bon nombre de colonnes
    table = doc.add_table(rows=1, cols=num_cols)

    # Ajouter les en-têtes
    headers = ['Matières']
    for periode in periodes:
        headers.append(f"{periode.nom} - Note Obtenue")
        headers.append(f"{periode.nom} - Note Attendue")
    headers.append('Total Obtenu')
    headers.append('Total Attendu')
    for periode in periodes:
        headers.append(f"{periode.nom} - Pourcentage")
    
    for i, header in enumerate(headers):
        table.cell(0, i).text = header

    # Variables pour accumuler les totaux
    total_obtenu_periode = [0] * len(periodes)
    total_attendu_periode = [0] * len(periodes)

    # Ajouter les données des notations
    for matiere in matieres:
        row = table.add_row().cells
        row[0].text = matiere.nom  # Commencer la ligne avec le nom de la matière
        for index, periode in enumerate(periodes):
            notation = Notation.objects.filter(eleve=eleve, matiere=matiere, periode=periode).first()
            if notation:
                row[index * 2 + 1].text = str(notation.note_obtenue)
                row[index * 2 + 2].text = str(notation.note_attendue)
                total_obtenu_periode[index] += notation.note_obtenue
                total_attendu_periode[index] += notation.note_attendue
            else:
                row[index * 2 + 1].text = 'N/A'
                row[index * 2 + 2].text = 'N/A'

    # Ajouter la ligne de totaux
    total_row = table.add_row().cells
    total_row[0].text = 'Total'
    for i, (total_obtenu, total_attendu) in enumerate(zip(total_obtenu_periode, total_attendu_periode)):
        total_row[i * 2 + 1].text = str(total_obtenu)
        total_row[i * 2 + 2].text = str(total_attendu)

    # Calculer et ajouter les pourcentages
    pourcentage_row = table.add_row().cells
    pourcentage_row[0].text = 'Pourcentage'
    for i, (total_obtenu, total_attendu) in enumerate(zip(total_obtenu_periode, total_attendu_periode)):
        if total_attendu > 0:
            pourcentage = (total_obtenu / total_attendu) * 100
        else:
            pourcentage = 0
        pourcentage_row[i * 2 + 1].text = f"{pourcentage:.2f}%"

    # Créer la réponse HTTP
    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = f'attachment; filename="{eleve.nom}_situation.docx"'

    # Sauvegarder le document dans la réponse
    doc.save(response)
    
    return response

# periode par eleve
from django.http import HttpResponse
from django.shortcuts import get_object_or_404
from .models import Eleve, Periode, Notation
from docx import Document
from django.db.models import Sum

def generer_word_periode(request, eleve_id, periode_id):
    # Récupérer les objets Eleve et Periode
    eleve = get_object_or_404(Eleve, id=eleve_id)
    periode = get_object_or_404(Periode, id=periode_id)
    
    # Récupérer les notations de l'élève pour la période
    notations = Notation.objects.filter(eleve=eleve, periode=periode)

    # Calcul des totaux obtenus et attendus
    total_attendu = notations.aggregate(Sum('note_attendue'))['note_attendue__sum'] or 0
    total_obtenu = notations.aggregate(Sum('note_obtenue'))['note_obtenue__sum'] or 0

    # Calcul du pourcentage
    pourcentage = (total_obtenu / total_attendu * 100) if total_attendu > 0 else 0

    # Créer un document Word
    doc = Document()

    # Ajouter un en-tête personnalisé
    doc.add_heading(f'Rapport de Note - {eleve.nom}', 0)
    
    # Ajouter des détails sur l'élève et la période
    doc.add_paragraph(f'Élève: {eleve.nom}')
    doc.add_paragraph(f'Classe: {eleve.classe.nom}')
    doc.add_paragraph(f'École: Mon École ABC')  # Nom de l'école, vous pouvez personnaliser cette valeur
    doc.add_paragraph(f'Période: {periode.nom}')
    
    # Ajouter un tableau pour les notations
    table = doc.add_table(rows=1, cols=4)
    
    # Ajouter les en-têtes de la table
    table.cell(0, 0).text = 'Matière'
    table.cell(0, 1).text = 'Note Attendue'
    table.cell(0, 2).text = 'Note Obtenue'
    table.cell(0, 3).text = 'Commentaire'  # Si vous n'avez pas de champ commentaire, vous pouvez laisser cela ou le personnaliser

    # Remplir les lignes du tableau avec les données de la notation
    for notation in notations:
        row = table.add_row().cells
        row[0].text = notation.matiere.nom
        row[1].text = str(notation.note_attendue)
        row[2].text = str(notation.note_obtenue)
        row[3].text = 'N/A'  # Pas de champ 'commentaire', donc on met 'N/A' par défaut

    # Ajouter le total et le pourcentage en bas du tableau
    doc.add_paragraph(f'Total Attendu: {total_attendu}')
    doc.add_paragraph(f'Total Obtenu: {total_obtenu}')
    doc.add_paragraph(f'Pourcentage: {pourcentage:.2f}%')

    # Créer la réponse HTTP pour permettre le téléchargement du fichier Word
    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = f'attachment; filename="{eleve.nom}_{periode.nom}_details.docx"'

    # Sauvegarder le document dans la réponse
    doc.save(response)

    return response

# liste de notes par matieres et par periode 
from django.http import HttpResponse
from django.shortcuts import get_object_or_404
from .models import Matiere, Periode, Notation
from docx import Document

def generer_word_periode_matiere(request, matiere_id, periode_id):
    # Récupérer les objets Matiere et Periode
    matiere = get_object_or_404(Matiere, id=matiere_id)
    periode = get_object_or_404(Periode, id=periode_id)
    
    # Récupérer les notations pour cette matière et cette période, avec les élèves
    notations = Notation.objects.filter(matiere=matiere, periode=periode).select_related('eleve')

    # Créer un document Word
    doc = Document()

    # Ajouter l'en-tête avec le nom de la matière, de la période, et du professeur
    doc.add_heading(f'Rapport de Notation - {matiere.nom}', 0)
    doc.add_paragraph(f'Matière : {matiere.nom}')
    doc.add_paragraph(f'Période : {periode.nom}')
    doc.add_paragraph(f'Professeur : {matiere.prof.username}')  # Utiliser le champ `username` pour le professeur
    
    # Ajouter un tableau pour les notations des élèves
    table = doc.add_table(rows=1, cols=4)
    
    # Ajouter les en-têtes de la table
    table.cell(0, 0).text = 'Élève'
    table.cell(0, 1).text = 'Note Attendue'
    table.cell(0, 2).text = 'Note Obtenue'
    table.cell(0, 3).text = 'Commentaire'  # Si vous avez des commentaires dans vos notations, vous pouvez ajouter une colonne pour ça

    # Remplir les lignes du tableau avec les données des notations
    for notation in notations:
        row = table.add_row().cells
        # Utiliser nom et prénom de l'élève
        row[0].text = f"{notation.eleve.nom} "  # Nom et Prénom
        row[1].text = str(notation.note_attendue)  # Note attendue
        row[2].text = str(notation.note_obtenue)  # Note obtenue
        row[3].text = 'N/A'  # Ici, il n'y a pas de champ commentaire dans Notation, donc on met "N/A"

    # Créer la réponse HTTP pour permettre le téléchargement du fichier Word
    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = f'attachment; filename="{matiere.nom}_{periode.nom}_rapport.docx"'

    # Sauvegarder le document dans la réponse
    doc.save(response)

    return response
